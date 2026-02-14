#!/usr/bin/env python3
"""
═══════════════════════════════════════════════════════════════════════════════
MS3 SIGMA DIVERGENCE DEFINITIVE PIPELINE — Google Colab (12GB RAM)
═══════════════════════════════════════════════════════════════════════════════

MANUSCRIPT: "Sigma-1 and Sigma-2 Receptors Exhibit Divergent Genome-Wide
Co-Expression Networks in Human Brain Despite Shared Subcellular
Localization"

TARGET:     Journal of Neurochemistry (JNC) — Wiley / International Society
            for Neurochemistry

AUTHOR:     Drake H. Harbert (ORCID: 0009-0007-7740-3616)
            Inner Architecture LLC, Canton, OH 44721, USA

VALIDATES:  Every number reported in the manuscript:
  - SIGMAR1/TMEM97 top 5% networks, thresholds, Jaccard, Fisher's, Spearman
  - 7-gene panel (21 pairwise comparisons) incl. SIGMAR1-LTN1
  - gProfiler enrichment (GO:BP, GO:MF, GO:CC, KEGG, Reactome)
  - 6 custom gene sets (MAM, sigma, UPR, methylation, vascular neg control, RQC)
  - Cell-type deconvolution (53 markers, 6 cell types)
  - Covariate adjustment (age, sex)
  - Multi-region replication (5 brain regions, cross-region ρ)
  - All 3 manuscript figures

GENERATES:  Supplementary Tables S1–S4 (.docx and .xlsx) for JNC submission

USAGE:      Copy cells into Google Colab. Run in order.
            Total runtime ~20-30 minutes depending on download speed.
            All outputs saved to Google Drive.

OUTPUTS:    All CSV results, all 3 manuscript figures, full enrichment data,
            Supplementary Tables S1–S4.
═══════════════════════════════════════════════════════════════════════════════
"""

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 1: ENVIRONMENT SETUP + GOOGLE DRIVE MOUNT                         ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

# ── Mount Google Drive ──
from google.colab import drive
drive.mount('/content/drive')

!pip install gprofiler-official requests matplotlib-venn python-docx openpyxl -q

import os
import gc
import json
import warnings
import numpy as np
import pandas as pd
from scipy import stats
from scipy.stats import pearsonr, spearmanr, fisher_exact
from collections import OrderedDict
from itertools import combinations
import matplotlib
matplotlib.rcParams['pdf.fonttype'] = 42
matplotlib.rcParams['ps.fonttype'] = 42
import matplotlib.pyplot as plt
from matplotlib_venn import venn2
warnings.filterwarnings('ignore')

# ── Output directories on Google Drive ──
DRIVE_BASE = "/content/drive/MyDrive/MS3_JNC_Submission"
RESULTS_DIR = os.path.join(DRIVE_BASE, "MS3_Results")
FIGURES_DIR = os.path.join(DRIVE_BASE, "MS3_Figures")
SUPPL_DIR   = os.path.join(DRIVE_BASE, "MS3_Supplementary_Tables")

os.makedirs(RESULTS_DIR, exist_ok=True)
os.makedirs(FIGURES_DIR, exist_ok=True)
os.makedirs(SUPPL_DIR, exist_ok=True)

print("=" * 70)
print("MS3 SIGMA DIVERGENCE DEFINITIVE PIPELINE")
print("Target: Journal of Neurochemistry (JNC)")
print("Drake H. Harbert — Inner Architecture LLC")
print("=" * 70)
print(f"\nDrive output: {DRIVE_BASE}")
print(f"  Results:    {RESULTS_DIR}")
print(f"  Figures:    {FIGURES_DIR}")
print(f"  Suppl.:     {SUPPL_DIR}")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 2: CONFIGURATION — LOCKED PARAMETERS                              ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

# ── Primary targets ──
PRIMARY_TARGETS = ['SIGMAR1', 'TMEM97']

# ── Full 7-gene panel (manuscript Section 2.3) ──
TARGET_GENES = ['EIF2S1', 'PELO', 'LTN1', 'NEMF', 'TMEM97', 'HSPA5', 'SIGMAR1']

# ── Brain regions (GTEx v8 exact SMTSD labels) ──
BRAIN_REGIONS = OrderedDict([
    ('BA9',          'Brain - Frontal Cortex (BA9)'),
    ('Putamen',      'Brain - Putamen (basal ganglia)'),
    ('Hippocampus',  'Brain - Hippocampus'),
    ('NAcc',         'Brain - Nucleus accumbens (basal ganglia)'),
    ('BA24',         'Brain - Anterior cingulate cortex (BA24)'),
])
PRIMARY_REGION = 'BA9'

# ── Expression filter ──
MIN_MEDIAN_TPM = 1.0

# ── Network threshold ──
TOP_PERCENT = 5  # top 5%

# ── Custom gene sets (LOCKED to manuscript Section 2.4) ──

# (1) MAM-mitochondrial markers: 14 genes
MAM_MITO_GENES = [
    'VDAC1', 'VDAC2', 'VDAC3', 'MFN1', 'MFN2', 'RHOT1', 'RHOT2',
    'ITPR1', 'ITPR2', 'ITPR3', 'VAPB', 'RMDN3', 'PACS2', 'FATE1',
]

# (2) Sigma receptor network: 4 genes
SIGMA_NETWORK_GENES = ['SIGMAR1', 'TMEM97', 'PGRMC1', 'NPC1']

# (3) ER stress/UPR: 14 genes
ER_STRESS_UPR_GENES = [
    'HSPA5', 'HSP90B1', 'DDIT3', 'ATF4', 'ATF6', 'ERN1', 'EIF2AK3',
    'XBP1', 'DNAJB9', 'HERPUD1', 'EDEM1', 'CALR', 'CANX', 'P4HB',
]

# (4) Methylation pathway: 18 genes
METHYLATION_GENES = [
    'MAT1A', 'MAT2A', 'MAT2B', 'AHCY', 'AHCYL1', 'AHCYL2',
    'MTR', 'MTRR', 'MTHFR', 'BHMT', 'BHMT2', 'CBS', 'CTH',
    'GNMT', 'PEMT', 'NNMT', 'INMT', 'DNMT1',
]

# (5) Vascular markers (NEGATIVE CONTROL): 18 genes
VASCULAR_GENES = [
    'PECAM1', 'CDH5', 'VWF', 'FLT1', 'KDR', 'ENG', 'CLDN5', 'ESAM',
    'ERG', 'TIE1', 'TEK', 'ANGPT1', 'ANGPT2', 'NOS3', 'MCAM',
    'PODXL', 'EMCN', 'ROBO4',
]

# (6) RQC genes (validates Section 3.5 claim: VCP, NPLOC4, TCF25)
RQC_GENES = [
    'PELO', 'HBS1L', 'LTN1', 'NEMF', 'ANKZF1',
    'VCP', 'UFD1', 'NPLOC4', 'ZNF598', 'RACK1', 'ABCE1', 'TCF25',
]

# ── Cell-type deconvolution markers (53 genes, 6 cell types) ──
# From Darmanis et al. 2015 + Lake et al. 2018
CELLTYPE_MARKERS = {
    'Neurons': [
        'SNAP25', 'SYT1', 'GAD1', 'GAD2', 'SLC17A7', 'RBFOX3',
        'STMN2', 'SYN1', 'NRGN',
    ],
    'Astrocytes': [
        'AQP4', 'GFAP', 'SLC1A2', 'SLC1A3', 'ALDH1L1', 'GJA1',
        'S100B', 'SOX9', 'GLUL',
    ],
    'Oligodendrocytes': [
        'MBP', 'MOG', 'PLP1', 'MAG', 'MOBP', 'CLDN11', 'CNP',
        'OPALIN', 'TF',
    ],
    'Microglia': [
        'CX3CR1', 'P2RY12', 'CSF1R', 'TMEM119', 'AIF1', 'ITGAM',
        'CD68', 'HEXB', 'TREM2',
    ],
    'Endothelial': [
        'CLDN5', 'FLT1', 'PECAM1', 'VWF', 'CDH5', 'ERG',
        'ESAM', 'TIE1',
    ],
    'OPCs': [
        'PDGFRA', 'CSPG4', 'OLIG1', 'OLIG2', 'SOX10', 'NKX2-2',
        'GPR17', 'PCDH15', 'NEU4',
    ],
}
ALL_CELLTYPE_MARKERS = [g for markers in CELLTYPE_MARKERS.values()
                         for g in markers]

# Assertions
assert len(MAM_MITO_GENES) == 14, f"MAM must be 14, got {len(MAM_MITO_GENES)}"
assert len(SIGMA_NETWORK_GENES) == 4, f"Sigma must be 4, got {len(SIGMA_NETWORK_GENES)}"
assert len(ER_STRESS_UPR_GENES) == 14, f"UPR must be 14, got {len(ER_STRESS_UPR_GENES)}"
assert len(METHYLATION_GENES) == 18, f"Methylation must be 18, got {len(METHYLATION_GENES)}"
assert len(VASCULAR_GENES) == 18, f"Vascular must be 18, got {len(VASCULAR_GENES)}"

print(f"Primary targets: {PRIMARY_TARGETS}")
print(f"7-gene panel: {TARGET_GENES}")
print(f"Brain regions: {list(BRAIN_REGIONS.keys())}")
print(f"Custom sets: MAM={len(MAM_MITO_GENES)}, Sigma={len(SIGMA_NETWORK_GENES)}, "
      f"UPR={len(ER_STRESS_UPR_GENES)}, Methyl={len(METHYLATION_GENES)}, "
      f"Vascular={len(VASCULAR_GENES)}, RQC={len(RQC_GENES)}")
print(f"Cell-type markers: {len(ALL_CELLTYPE_MARKERS)} across "
      f"{len(CELLTYPE_MARKERS)} types")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 3: DOWNLOAD GTEx v8 DATA                                          ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

GTEX_TPM_URL = "https://storage.googleapis.com/adult-gtex/bulk-gex/v8/rna-seq/GTEx_Analysis_2017-06-05_v8_RNASeQCv1.1.9_gene_tpm.gct.gz"
GTEX_SAMPLE_URL = "https://storage.googleapis.com/adult-gtex/annotations/v8/metadata-files/GTEx_Analysis_v8_Annotations_SampleAttributesDS.txt"
GTEX_SUBJECT_URL = "https://storage.googleapis.com/adult-gtex/annotations/v8/metadata-files/GTEx_Analysis_v8_Annotations_SubjectPhenotypesDS.txt"

TPM_FILE = "GTEx_v8_tpm.gct.gz"
SAMPLE_FILE = "GTEx_v8_sample_attributes.txt"
SUBJECT_FILE = "GTEx_v8_subject_phenotypes.txt"

import subprocess

for url, fname in [(GTEX_TPM_URL, TPM_FILE),
                   (GTEX_SAMPLE_URL, SAMPLE_FILE),
                   (GTEX_SUBJECT_URL, SUBJECT_FILE)]:
    if not os.path.exists(fname):
        print(f"Downloading {fname}...")
        subprocess.run(['wget', '-q', '-O', fname, url], check=True)
        print(f"  ✓ Downloaded {fname}")
    else:
        print(f"  ✓ {fname} exists")

print("\nAll GTEx files ready.")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 4: LOAD SAMPLE METADATA + SUBJECT PHENOTYPES                      ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print("Loading sample metadata...")
sample_attr = pd.read_csv(SAMPLE_FILE, sep='\t', low_memory=False)

# ── Diagnostic: Print actual GTEx brain tissue labels ──
print("\n  ── Actual GTEx brain tissue labels in SMTSD column ──")
brain_labels = sample_attr[
    sample_attr['SMTSD'].str.contains('Brain', na=False)
]['SMTSD'].unique()
for label in sorted(brain_labels):
    count = (sample_attr['SMTSD'] == label).sum()
    marker = " ◄ USED" if label in BRAIN_REGIONS.values() else ""
    print(f"    {repr(label):60s} n={count}{marker}")

# Verify all requested regions were found
for key, smtsd in BRAIN_REGIONS.items():
    if smtsd not in brain_labels:
        print(f"\n  ⚠ FATAL: '{smtsd}' NOT FOUND in GTEx labels!")
        raise ValueError(f"Region label mismatch for {key}: '{smtsd}'")

# Build region → sample ID mapping
print()
region_samples = {}
for key, smtsd in BRAIN_REGIONS.items():
    mask = sample_attr['SMTSD'] == smtsd
    region_samples[key] = list(sample_attr.loc[mask, 'SAMPID'])
    print(f"  {key}: {len(region_samples[key])} samples")

# Sanity check
for key, samples in region_samples.items():
    if len(samples) == 0:
        raise ValueError(f"Region {key} has 0 samples!")

# Load subject phenotypes (for covariate adjustment)
print("\nLoading subject phenotypes...")
subject_pheno = pd.read_csv(SUBJECT_FILE, sep='\t')
print(f"  Subjects: {len(subject_pheno)}")

def sample_to_subject(sample_id):
    parts = sample_id.split('-')
    if len(parts) >= 2:
        return '-'.join(parts[:2])
    return sample_id

# Build covariate frame for BA9
ba9_samples_list = region_samples[PRIMARY_REGION]
ba9_subjects = [sample_to_subject(s) for s in ba9_samples_list]
ba9_subject_df = pd.DataFrame({'SAMPID': ba9_samples_list, 'SUBJID': ba9_subjects})
ba9_subject_df = ba9_subject_df.merge(subject_pheno, on='SUBJID', how='left')

print(f"\nBA9 covariates:")
if 'DTHHRDY' in ba9_subject_df.columns:
    print(f"  Hardy Scale:\n{ba9_subject_df['DTHHRDY'].value_counts().sort_index()}")
if 'AGE' in ba9_subject_df.columns:
    print(f"  Age bins:\n{ba9_subject_df['AGE'].value_counts().sort_index()}")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 5: LOAD GTEx TPM — MEMORY-EFFICIENT (12GB SAFE)                   ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print("Loading GTEx TPM (brain regions only, memory-efficient)...")

import gzip

# Step 1: Read header to map column positions
print("  Reading header...")
with gzip.open(TPM_FILE, 'rt') as f:
    f.readline()  # version
    f.readline()  # dimensions
    header = f.readline().strip().split('\t')

all_brain_samples = set()
for key in BRAIN_REGIONS:
    all_brain_samples.update(region_samples[key])

sample_cols = {}
for i, col in enumerate(header):
    if col in all_brain_samples:
        sample_cols[i] = col

keep_cols = [0, 1] + sorted(sample_cols.keys())
print(f"  Brain samples found: {len(sample_cols)} / {len(all_brain_samples)}")
print(f"  Total columns in file: {len(header)}")
print(f"  Columns to load: {len(keep_cols)}")

# Step 2: Stream-read, keeping only brain columns
print("  Loading expression data (this takes 3-5 minutes)...")
chunks = []
chunk_size = 5000
with gzip.open(TPM_FILE, 'rt') as f:
    f.readline()
    f.readline()
    f.readline()

    row_buffer = []
    for line_num, line in enumerate(f):
        parts = line.strip().split('\t')
        row = [parts[i] if i < len(parts) else '' for i in keep_cols]
        row_buffer.append(row)

        if len(row_buffer) >= chunk_size:
            chunks.append(pd.DataFrame(row_buffer))
            row_buffer = []
            if line_num % 10000 == 0:
                print(f"    Processed {line_num:,} genes...")

    if row_buffer:
        chunks.append(pd.DataFrame(row_buffer))

print("  Concatenating chunks...")
tpm_raw = pd.concat(chunks, ignore_index=True)
del chunks, row_buffer
gc.collect()

col_names = ['Name', 'Description'] + [sample_cols[i] for i in sorted(sample_cols.keys())]
tpm_raw.columns = col_names
tpm_raw = tpm_raw.set_index('Name')
gene_descriptions = tpm_raw['Description'].copy()
tpm_raw = tpm_raw.drop('Description', axis=1)
tpm_raw = tpm_raw.apply(pd.to_numeric, errors='coerce').astype(np.float32)

print(f"\n  Loaded: {tpm_raw.shape[0]} genes × {tpm_raw.shape[1]} samples")
print(f"  Memory: {tpm_raw.memory_usage(deep=True).sum() / 1e9:.2f} GB")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 6: EXTRACT REGION-SPECIFIC MATRICES + GENE FILTERING              ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

def prepare_region(region_key):
    """Extract and filter expression matrix for one brain region."""
    samples = [s for s in region_samples[region_key] if s in tpm_raw.columns]
    expr = tpm_raw[samples].copy()
    expr['gene_symbol'] = gene_descriptions.reindex(expr.index)
    expr = expr.dropna(subset=['gene_symbol'])
    expr = expr[expr['gene_symbol'] != '']

    # For duplicate gene symbols, keep highest median
    expr['median_tpm'] = expr[samples].median(axis=1)
    expr = expr.sort_values('median_tpm', ascending=False)
    expr = expr.drop_duplicates(subset='gene_symbol', keep='first')

    # Filter: median TPM >= 1.0
    expr = expr[expr['median_tpm'] >= MIN_MEDIAN_TPM]

    expr = expr.set_index('gene_symbol')
    expr = expr.drop('median_tpm', axis=1)

    # Log2 transform
    log2_expr = np.log2(expr + 1)

    return log2_expr, len(samples), len(log2_expr), list(log2_expr.index)

print("Preparing brain region matrices...\n")
region_data = {}
for key in BRAIN_REGIONS:
    log2_expr, n_samp, n_genes, genes = prepare_region(key)
    region_data[key] = {
        'expr': log2_expr,
        'n_samples': n_samp,
        'n_genes': n_genes,
        'genes': genes,
    }
    present = [g for g in TARGET_GENES if g in genes]
    missing = [g for g in TARGET_GENES if g not in genes]
    print(f"  {key}: n={n_samp}, genes={n_genes}, targets={len(present)}/7"
          + (f" (missing: {missing})" if missing else ""))

    # Report SIGMAR1 median TPM
    if 'SIGMAR1' in genes:
        sigmar1_tpm = tpm_raw.loc[
            gene_descriptions[gene_descriptions == 'SIGMAR1'].index
        ]
        if len(sigmar1_tpm) > 0:
            samples_r = [s for s in region_samples[key] if s in sigmar1_tpm.columns]
            if samples_r:
                median_val = sigmar1_tpm[samples_r].median(axis=1).values[0]
                if key == PRIMARY_REGION:
                    print(f"    SIGMAR1 median TPM = {median_val:.2f}")

gc.collect()

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 7: GENOME-WIDE CORRELATIONS (PRIMARY REGION — ALL 7 TARGETS)      ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print(f"GENOME-WIDE CO-EXPRESSION ANALYSIS — {PRIMARY_REGION}")
print(f"{'='*70}\n")

ba9 = region_data[PRIMARY_REGION]['expr']
n_genes_ba9 = region_data[PRIMARY_REGION]['n_genes']
n_samples_ba9 = region_data[PRIMARY_REGION]['n_samples']

print(f"Matrix: {n_genes_ba9} genes × {n_samples_ba9} samples")

# Vectorized Pearson correlation for each target
correlations = {}
for target in TARGET_GENES:
    if target not in ba9.index:
        print(f"  ⚠ {target} not found in {PRIMARY_REGION}")
        continue

    target_expr = ba9.loc[target].values.astype(np.float64)
    other_genes = [g for g in ba9.index if g != target]
    other_expr = ba9.loc[other_genes].values.astype(np.float64)

    x = target_expr - target_expr.mean()
    Y = other_expr - other_expr.mean(axis=1, keepdims=True)
    x_std = np.sqrt(np.sum(x**2))
    Y_std = np.sqrt(np.sum(Y**2, axis=1))

    valid = Y_std > 0
    r_values = np.full(len(other_genes), np.nan)
    r_values[valid] = np.dot(Y[valid], x) / (Y_std[valid] * x_std)

    corr_series = pd.Series(r_values, index=other_genes, name=target)
    correlations[target] = corr_series.dropna().sort_values(ascending=False)

    print(f"  {target}: {len(corr_series.dropna())} genes, "
          f"top = {corr_series.dropna().idxmax()} "
          f"(r = {corr_series.dropna().max():.3f})")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 8: TOP 5% NETWORKS + ALL 21 PAIRWISE COMPARISONS                  ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("TOP 5% NETWORK EXTRACTION + PAIRWISE COMPARISONS")
print(f"{'='*70}\n")

networks = {}
thresholds = {}

for target, corr in correlations.items():
    n_total = len(corr)
    n_top = int(np.ceil(n_total * TOP_PERCENT / 100))
    top_genes = corr.head(n_top)
    threshold = top_genes.iloc[-1]
    networks[target] = set(top_genes.index)
    thresholds[target] = threshold
    print(f"  {target}: top 5% = {n_top} genes, r ≥ {threshold:.3f}")

gene_universe = len(correlations[TARGET_GENES[0]])
print(f"\n  Gene universe (non-target): {gene_universe}")

# ── All 21 pairwise comparisons ──
pairwise_results = []
for g1, g2 in combinations(TARGET_GENES, 2):
    if g1 not in networks or g2 not in networks:
        continue

    set1, set2 = networks[g1], networks[g2]
    shared = set1 & set2
    union = set1 | set2
    only1 = set1 - set2
    only2 = set2 - set1

    jaccard = len(shared) / len(union) if len(union) > 0 else 0

    a = len(shared)
    b = len(only1)
    c = len(only2)
    d = gene_universe - len(union)
    table = np.array([[a, b], [c, d]])
    fisher_or, fisher_p = fisher_exact(table, alternative='greater')

    common_genes = sorted(set(correlations[g1].index) & set(correlations[g2].index))
    ranks1 = correlations[g1].reindex(common_genes).rank(ascending=False)
    ranks2 = correlations[g2].reindex(common_genes).rank(ascending=False)
    rho, _ = spearmanr(ranks1, ranks2)

    pairwise_results.append({
        'gene1': g1, 'gene2': g2,
        'shared': len(shared), 'jaccard': jaccard,
        'fisher_or': fisher_or, 'fisher_p': fisher_p,
        'spearman_rho': rho,
        'set1_size': len(set1), 'set2_size': len(set2),
    })

comp_df = pd.DataFrame(pairwise_results).sort_values('jaccard', ascending=False)
comp_df.to_csv(os.path.join(RESULTS_DIR, "all_21_pairwise_comparisons.csv"), index=False)

# Print SIGMAR1 and TMEM97 comparisons
print("\n─── SIGMAR1 pairwise ───")
for _, row in comp_df[
    (comp_df['gene1'] == 'SIGMAR1') | (comp_df['gene2'] == 'SIGMAR1')
].iterrows():
    g1, g2 = row['gene1'], row['gene2']
    print(f"  {g1}–{g2}: J={row['jaccard']:.3f}, shared={int(row['shared'])}, "
          f"OR={row['fisher_or']:.2f}, p={row['fisher_p']:.2e}, ρ={row['spearman_rho']:.3f}")

print("\n─── TMEM97 pairwise ───")
for _, row in comp_df[
    (comp_df['gene1'] == 'TMEM97') | (comp_df['gene2'] == 'TMEM97')
].iterrows():
    g1, g2 = row['gene1'], row['gene2']
    print(f"  {g1}–{g2}: J={row['jaccard']:.3f}, shared={int(row['shared'])}, "
          f"OR={row['fisher_or']:.2f}, p={row['fisher_p']:.2e}, ρ={row['spearman_rho']:.3f}")

# ── KEY MANUSCRIPT VALIDATION: SIGMAR1–TMEM97 ──
st_row = comp_df[
    ((comp_df['gene1'] == 'SIGMAR1') & (comp_df['gene2'] == 'TMEM97')) |
    ((comp_df['gene1'] == 'TMEM97') & (comp_df['gene2'] == 'SIGMAR1'))
]
print(f"\n{'─'*50}")
print("MANUSCRIPT VALIDATION: SIGMAR1–TMEM97")
print(f"  Shared genes: {int(st_row['shared'].values[0])}")
print(f"  Jaccard: {st_row['jaccard'].values[0]:.3f}")
print(f"  Fisher OR: {st_row['fisher_or'].values[0]:.2f}")
print(f"  Fisher p: {st_row['fisher_p'].values[0]:.2e}")
print(f"  Spearman ρ: {st_row['spearman_rho'].values[0]:.3f}")
print(f"{'─'*50}")

# ── SIGMAR1–LTN1 overlap check ──
sl_row = comp_df[
    ((comp_df['gene1'] == 'SIGMAR1') & (comp_df['gene2'] == 'LTN1')) |
    ((comp_df['gene1'] == 'LTN1') & (comp_df['gene2'] == 'SIGMAR1'))
]
if len(sl_row) > 0:
    print(f"\nSIGMAR1–LTN1: shared={int(sl_row['shared'].values[0])}, "
          f"J={sl_row['jaccard'].values[0]:.3f}, "
          f"ρ={sl_row['spearman_rho'].values[0]:.3f}")

# Verification: all Jaccard internally consistent
for _, row in comp_df.iterrows():
    s, s1, s2 = int(row['shared']), int(row['set1_size']), int(row['set2_size'])
    expected_j = s / (s1 + s2 - s) if (s1 + s2 - s) > 0 else 0
    assert abs(row['jaccard'] - expected_j) < 0.001
print("\n  ✓ All Jaccard indices internally consistent")

# ── SIGMAR1/TMEM97 unique and shared gene sets ──
sigmar1_unique = sorted(networks['SIGMAR1'] - networks['TMEM97'])
tmem97_unique = sorted(networks['TMEM97'] - networks['SIGMAR1'])
shared_st = sorted(networks['SIGMAR1'] & networks['TMEM97'])

print(f"\n  SIGMAR1 unique: {len(sigmar1_unique)} genes")
print(f"  TMEM97 unique:  {len(tmem97_unique)} genes")
print(f"  Shared:         {len(shared_st)} genes")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 9: TOP CO-EXPRESSION PARTNERS                                      ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("TOP CO-EXPRESSION PARTNERS")
print(f"{'='*70}\n")

for target in PRIMARY_TARGETS:
    print(f"\n  {target} top 10:")
    top10 = correlations[target].head(10)
    for rank, (gene, r_val) in enumerate(top10.items(), 1):
        print(f"    {rank:2d}. {gene:12s} r = {r_val:.3f}")

# Validate manuscript claims for SIGMAR1 top 5
print(f"\n─── SIGMAR1 top 5 validation ───")
ms_top5 = ['YIPF3', 'RAB1B', 'AAMP', 'PSMD3', 'ARF1']
for gene in ms_top5:
    if gene in correlations['SIGMAR1'].index:
        r = correlations['SIGMAR1'][gene]
        rank = (correlations['SIGMAR1'] > r).sum() + 1
        print(f"  {gene}: r = {r:.3f}, rank = {rank}")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 10: CUSTOM GENE SET ENRICHMENT (Fisher's exact)                    ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("CUSTOM GENE SET ENRICHMENT")
print(f"{'='*70}\n")

def compute_custom_enrichment(network_genes, gene_set_list, set_name,
                               all_genes_list, verbose=True):
    """Fisher's exact test for custom gene set enrichment."""
    universe = set(all_genes_list)
    expressed = [g for g in gene_set_list if g in universe]
    in_network = [g for g in expressed if g in network_genes]

    n_expressed = len(expressed)
    n_in_network = len(in_network)
    n_network = len(network_genes)
    n_universe = len(universe)

    if n_expressed == 0:
        return None

    expected = n_expressed * n_network / n_universe
    fold = n_in_network / expected if expected > 0 else 0

    a = n_in_network
    b = n_network - n_in_network
    c = n_expressed - n_in_network
    d = n_universe - n_network - c
    table = np.array([[a, b], [c, d]])
    _, p_val = fisher_exact(table, alternative='greater')

    result = {
        'gene_set': set_name,
        'expressed': n_expressed,
        'in_network': n_in_network,
        'fold_enrichment': fold,
        'p_value': p_val,
        'genes_found': ', '.join(sorted(in_network)),
    }

    if verbose:
        sig = "✓" if p_val < 0.05 else "✗"
        print(f"  {sig} {set_name}: {n_in_network}/{n_expressed}, "
              f"{fold:.1f}×, p = {p_val:.2e}")
        if in_network:
            print(f"    Genes: {', '.join(sorted(in_network))}")

    return result

custom_sets = OrderedDict([
    ('MAM-mitochondrial', MAM_MITO_GENES),
    ('Sigma receptor network', SIGMA_NETWORK_GENES),
    ('ER stress/UPR', ER_STRESS_UPR_GENES),
    ('Methylation pathway', METHYLATION_GENES),
    ('Vascular markers (neg ctrl)', VASCULAR_GENES),
    ('Ribosome quality control', RQC_GENES),
])

# Run for SIGMAR1 and TMEM97 — store results for supplementary table
all_custom_results = {}
for target in PRIMARY_TARGETS:
    print(f"\n─── {target} top 5% custom enrichment ───\n")
    all_genes_t = list(correlations[target].index) + [target]
    target_results = []
    for set_name, gene_list in custom_sets.items():
        result = compute_custom_enrichment(
            networks[target], gene_list, set_name, all_genes_t)
        if result:
            result['target'] = target
            target_results.append(result)

    all_custom_results[target] = target_results
    pd.DataFrame(target_results).to_csv(
        os.path.join(RESULTS_DIR, f"{target}_custom_enrichment.csv"), index=False)

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 11: gProfiler ENRICHMENT                                          ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("gProfiler GO/KEGG/REACTOME ENRICHMENT")
print(f"{'='*70}\n")

try:
    from gprofiler import GProfiler
    gp = GProfiler(return_dataframe=True)
    GPROFILER_AVAILABLE = True
    print("✓ gprofiler-official loaded")
except ImportError:
    GPROFILER_AVAILABLE = False
    print("⚠ gprofiler not available, using requests API fallback")

def run_gprofiler_enrichment(gene_list, background_list, label=""):
    """Run gProfiler enrichment with proper background."""
    genes = list(gene_list)
    bg = list(background_list)

    if GPROFILER_AVAILABLE:
        try:
            df = gp.profile(
                organism='hsapiens',
                query=genes,
                background=bg,
                sources=['GO:BP', 'GO:MF', 'GO:CC', 'KEGG', 'REAC'],
                significance_threshold_method='g_SCS',
                user_threshold=0.05,
                no_evidences=False,
            )
            if df is not None and len(df) > 0:
                print(f"  {label}: {len(df)} significant terms")
                return df
            else:
                print(f"  {label}: 0 significant terms")
                return pd.DataFrame()
        except Exception as e:
            print(f"  {label}: gProfiler error — {e}")
            return pd.DataFrame()
    else:
        import requests
        url = "https://biit.cs.ut.ee/gprofiler/api/gost/profile/"
        payload = {
            'organism': 'hsapiens',
            'query': genes,
            'sources': ['GO:BP', 'GO:MF', 'GO:CC', 'KEGG', 'REAC'],
            'user_threshold': 0.05,
            'significance_threshold_method': 'g_SCS',
            'no_evidences': True,
            'background': bg,
        }
        try:
            resp = requests.post(url, json=payload, timeout=120)
            if resp.status_code == 200:
                data = resp.json()
                results = data.get('result', [])
                if results:
                    df = pd.DataFrame(results)
                    print(f"  {label}: {len(df)} significant terms")
                    return df
                else:
                    print(f"  {label}: 0 terms")
                    return pd.DataFrame()
            else:
                print(f"  {label}: HTTP {resp.status_code}")
                return pd.DataFrame()
        except Exception as e:
            print(f"  {label}: API error — {e}")
            return pd.DataFrame()

# Background: all expressed genes
background = list(correlations['SIGMAR1'].index) + ['SIGMAR1']

# 5 enrichment runs matching manuscript Section 2.4
print("Running enrichment analyses:\n")
go_sigmar1_full = run_gprofiler_enrichment(
    list(networks['SIGMAR1']), background, "SIGMAR1 full")
go_tmem97_full = run_gprofiler_enrichment(
    list(networks['TMEM97']), background, "TMEM97 full")
go_sigmar1_unique = run_gprofiler_enrichment(
    sigmar1_unique, background, f"SIGMAR1 unique ({len(sigmar1_unique)} genes)")
go_tmem97_unique = run_gprofiler_enrichment(
    tmem97_unique, background, f"TMEM97 unique ({len(tmem97_unique)} genes)")
go_shared = run_gprofiler_enrichment(
    shared_st, background, f"Shared ({len(shared_st)} genes)")

# Save all results
all_gprofiler = {}
for name, df in [
    ('SIGMAR1_full', go_sigmar1_full),
    ('TMEM97_full', go_tmem97_full),
    ('SIGMAR1_unique', go_sigmar1_unique),
    ('TMEM97_unique', go_tmem97_unique),
    ('shared_SIGMAR1_TMEM97', go_shared),
]:
    all_gprofiler[name] = df
    if len(df) > 0:
        df.to_csv(os.path.join(RESULTS_DIR, f"gProfiler_{name}.csv"), index=False)

# Print top terms for validation
def print_top_terms(go_df, label, sources=['GO:BP', 'KEGG', 'REAC'], n=5):
    if go_df is None or len(go_df) == 0:
        return
    for source in sources:
        subset = go_df[go_df['source'] == source].sort_values('p_value')
        if len(subset) > 0:
            print(f"\n  {label} — {source} ({len(subset)} terms):")
            for _, row in subset.head(n).iterrows():
                print(f"    {row['name'][:55]}: p = {row['p_value']:.2e}")

print_top_terms(go_sigmar1_full, "SIGMAR1 full")
print_top_terms(go_tmem97_full, "TMEM97 full")
print_top_terms(go_sigmar1_unique, "SIGMAR1 unique")
print_top_terms(go_tmem97_unique, "TMEM97 unique")
print_top_terms(go_shared, "Shared")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 12: CELL-TYPE DECONVOLUTION SENSITIVITY                           ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("CELL-TYPE DECONVOLUTION SENSITIVITY ANALYSIS")
print(f"{'='*70}\n")

from numpy.linalg import lstsq

# Estimate cell-type proportions per sample
celltype_proportions = pd.DataFrame(index=ba9.columns)
for ct_name, markers in CELLTYPE_MARKERS.items():
    present = [g for g in markers if g in ba9.index]
    if present:
        celltype_proportions[ct_name] = ba9.loc[present].mean(axis=0)
        print(f"  {ct_name}: {len(present)}/{len(markers)} markers present")

print(f"\n  Cell-type proportion matrix: {celltype_proportions.shape}")

# Partial correlations controlling for cell-type composition
def partial_correlation_network(target, expr_df, covar_df, top_pct=5):
    """Compute genome-wide correlations after regressing out covariates."""
    valid = covar_df.dropna().index
    valid = [s for s in valid if s in expr_df.columns]
    n_valid = len(valid)

    if n_valid < 30:
        return None, None, None, n_valid

    expr_sub = expr_df[valid]
    covar_matrix = covar_df.loc[valid].values

    X = np.column_stack([np.ones(n_valid), covar_matrix])

    target_vals = expr_sub.loc[target].values.astype(np.float64)
    beta, _, _, _ = lstsq(X, target_vals, rcond=None)
    target_resid = target_vals - X @ beta

    other_genes = [g for g in expr_sub.index if g != target]
    other_vals = expr_sub.loc[other_genes].values.astype(np.float64)

    betas = lstsq(X, other_vals.T, rcond=None)[0]
    resid = other_vals - (X @ betas).T

    x = target_resid - target_resid.mean()
    Y = resid - resid.mean(axis=1, keepdims=True)
    x_std = np.sqrt(np.sum(x**2))
    Y_std = np.sqrt(np.sum(Y**2, axis=1))
    valid_mask = Y_std > 0
    r_vals = np.full(len(other_genes), np.nan)
    r_vals[valid_mask] = np.dot(Y[valid_mask], x) / (Y_std[valid_mask] * x_std)

    corr_series = pd.Series(r_vals, index=other_genes).dropna().sort_values(ascending=False)
    n_top = int(np.ceil(len(corr_series) * top_pct / 100))
    top5_set = set(corr_series.head(n_top).index)
    threshold = corr_series.iloc[n_top - 1] if n_top <= len(corr_series) else np.nan

    return corr_series, threshold, top5_set, n_valid

print("\nCell-type adjusted correlations:")
ct_results = {}
for target in PRIMARY_TARGETS:
    corr_ct, thr_ct, net_ct, n_valid = partial_correlation_network(
        target, ba9, celltype_proportions)

    if corr_ct is None:
        print(f"  ⚠ {target}: insufficient samples")
        continue

    ct_results[target] = {
        'corr': corr_ct, 'threshold': thr_ct, 'network': net_ct,
    }

    # Rank preservation
    common = sorted(set(corr_ct.index) & set(correlations[target].index))
    rho_preserve, _ = spearmanr(
        corr_ct.reindex(common).values,
        correlations[target].reindex(common).values)

    print(f"  {target}: threshold r ≥ {thr_ct:.3f} (raw: {thresholds[target]:.3f}), "
          f"rank preservation ρ = {rho_preserve:.4f}")

    # Check VCP still in top 5% for SIGMAR1
    if target == 'SIGMAR1' and net_ct:
        vcp_in = 'VCP' in net_ct
        print(f"    VCP in cell-type adjusted top 5%: {vcp_in}")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 13: COVARIATE ADJUSTMENT (AGE + SEX)                              ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("COVARIATE ADJUSTMENT: AGE + SEX")
print(f"{'='*70}\n")

# Build covariate matrix
ba9_covar = ba9_subject_df.set_index('SAMPID').reindex(ba9.columns)

age_map = {'20-29': 25, '30-39': 35, '40-49': 45, '50-59': 55,
           '60-69': 65, '70-79': 75}
ba9_covar['AGE_MID'] = ba9_covar['AGE'].map(age_map) if 'AGE' in ba9_covar.columns else np.nan
ba9_covar['SEX_NUM'] = ba9_covar['SEX'].astype(float) if 'SEX' in ba9_covar.columns else np.nan

covar_cols_agesex = ['AGE_MID', 'SEX_NUM']
covar_df_agesex = ba9_covar[covar_cols_agesex]

print(f"  Valid samples (age+sex): {covar_df_agesex.dropna().shape[0]}")

agesex_results = {}
for target in PRIMARY_TARGETS:
    corr_as, thr_as, net_as, n_valid = partial_correlation_network(
        target, ba9, covar_df_agesex)

    if corr_as is None:
        continue

    agesex_results[target] = {
        'corr': corr_as, 'threshold': thr_as, 'network': net_as,
    }

    common = sorted(set(corr_as.index) & set(correlations[target].index))
    rho_preserve, _ = spearmanr(
        corr_as.reindex(common).values,
        correlations[target].reindex(common).values)

    print(f"  {target}: threshold r ≥ {thr_as:.3f} (raw: {thresholds[target]:.3f}), "
          f"rank preservation ρ = {rho_preserve:.4f}")

# SIGMAR1-TMEM97 divergence after age+sex adjustment
if 'SIGMAR1' in agesex_results and 'TMEM97' in agesex_results:
    shared_adj = agesex_results['SIGMAR1']['network'] & agesex_results['TMEM97']['network']
    union_adj = agesex_results['SIGMAR1']['network'] | agesex_results['TMEM97']['network']
    j_adj = len(shared_adj) / len(union_adj) if len(union_adj) > 0 else 0
    print(f"\n  SIGMAR1-TMEM97 Jaccard (age+sex adjusted): {j_adj:.3f} "
          f"(raw: {st_row['jaccard'].values[0]:.3f})")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 14: MULTI-REGION REPLICATION                                       ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("MULTI-REGION REPLICATION")
print(f"{'='*70}\n")

# Compute genome-wide correlations for SIGMAR1 and TMEM97 in all regions
region_correlations = {}
for region_key in BRAIN_REGIONS:
    expr = region_data[region_key]['expr']
    n_samp = region_data[region_key]['n_samples']
    region_correlations[region_key] = {}

    if n_samp == 0:
        print(f"  ⚠ {region_key}: SKIPPED")
        continue

    for target in PRIMARY_TARGETS:
        if target not in expr.index:
            print(f"  ⚠ {target} not in {region_key}")
            continue

        target_vals = expr.loc[target].values.astype(np.float64)
        other_genes = [g for g in expr.index if g != target]
        other_vals = expr.loc[other_genes].values.astype(np.float64)

        x = target_vals - target_vals.mean()
        Y = other_vals - other_vals.mean(axis=1, keepdims=True)
        x_std = np.sqrt(np.sum(x**2))
        Y_std = np.sqrt(np.sum(Y**2, axis=1))
        valid = Y_std > 0
        r_vals = np.full(len(other_genes), np.nan)
        r_vals[valid] = np.dot(Y[valid], x) / (Y_std[valid] * x_std)

        corr_series = pd.Series(r_vals, index=other_genes).dropna().sort_values(ascending=False)
        region_correlations[region_key][target] = corr_series

    targets_done = list(region_correlations[region_key].keys())
    print(f"  {region_key} (n={n_samp}): {', '.join(targets_done)}")

# Cross-region Spearman rank correlations
print("\nCross-region rank correlations:")
region_keys = list(BRAIN_REGIONS.keys())
cross_region_matrix = {}

for target in PRIMARY_TARGETS:
    cross_region_matrix[target] = np.zeros((5, 5))
    for i, r1 in enumerate(region_keys):
        for j, r2 in enumerate(region_keys):
            if i == j:
                cross_region_matrix[target][i, j] = 1.0
                continue
            if (target in region_correlations.get(r1, {}) and
                target in region_correlations.get(r2, {})):
                corr1 = region_correlations[r1][target]
                corr2 = region_correlations[r2][target]
                common = sorted(set(corr1.index) & set(corr2.index))
                if len(common) > 100:
                    rho, _ = spearmanr(corr1.reindex(common).values,
                                       corr2.reindex(common).values)
                    cross_region_matrix[target][i, j] = rho

    print(f"\n  {target}:")
    for i, r1 in enumerate(region_keys):
        for j, r2 in enumerate(region_keys):
            if j > i:
                rho = cross_region_matrix[target][i, j]
                print(f"    {r1} vs {r2}: ρ = {rho:.3f}")

    vals = cross_region_matrix[target][np.triu_indices(5, k=1)]
    print(f"  Range: ρ = {vals.min():.3f} – {vals.max():.3f}")

# Multi-region custom enrichment for RQC (validates Section 3.5)
print("\n─── RQC enrichment across regions (validates Section 3.5) ───")
for region_key in BRAIN_REGIONS:
    if 'SIGMAR1' not in region_correlations.get(region_key, {}):
        continue
    corr = region_correlations[region_key]['SIGMAR1']
    n_top = int(np.ceil(len(corr) * TOP_PERCENT / 100))
    top5 = set(corr.head(n_top).index)
    all_genes_r = list(corr.index) + ['SIGMAR1']
    result = compute_custom_enrichment(top5, RQC_GENES, f"RQC in {region_key}",
                                        all_genes_r, verbose=False)
    if result:
        print(f"  {region_key}: {result['in_network']}/{result['expressed']}, "
              f"{result['fold_enrichment']:.1f}×, p = {result['p_value']:.2e}"
              f"  [{result['genes_found']}]")

# SIGMAR1-TMEM97 Jaccard across all regions
print("\n─── SIGMAR1–TMEM97 divergence across regions ───")
for region_key in BRAIN_REGIONS:
    if ('SIGMAR1' not in region_correlations.get(region_key, {}) or
        'TMEM97' not in region_correlations.get(region_key, {})):
        continue
    s_corr = region_correlations[region_key]['SIGMAR1']
    t_corr = region_correlations[region_key]['TMEM97']
    n_s = int(np.ceil(len(s_corr) * TOP_PERCENT / 100))
    n_t = int(np.ceil(len(t_corr) * TOP_PERCENT / 100))
    s_set = set(s_corr.head(n_s).index)
    t_set = set(t_corr.head(n_t).index)
    shared_r = s_set & t_set
    union_r = s_set | t_set
    j_r = len(shared_r) / len(union_r) if len(union_r) > 0 else 0
    print(f"  {region_key}: J = {j_r:.3f}, shared = {len(shared_r)}")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 15: EXPORT GENE LISTS                                              ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("EXPORTING GENE LISTS")
print(f"{'='*70}\n")

# Full top 5% for each primary target
for target in PRIMARY_TARGETS:
    top5_df = pd.DataFrame({
        'gene': correlations[target].head(len(networks[target])).index,
        f'r_with_{target}': correlations[target].head(len(networks[target])).values,
    })
    top5_df.to_csv(os.path.join(RESULTS_DIR, f"{target}_top5pct.csv"), index=False)
    print(f"  {target} top 5%: {len(top5_df)} genes")

# Shared, unique
pd.DataFrame({'gene': shared_st}).to_csv(
    os.path.join(RESULTS_DIR, "SIGMAR1_TMEM97_shared.csv"), index=False)
pd.DataFrame({'gene': sigmar1_unique}).to_csv(
    os.path.join(RESULTS_DIR, "SIGMAR1_unique.csv"), index=False)
pd.DataFrame({'gene': tmem97_unique}).to_csv(
    os.path.join(RESULTS_DIR, "TMEM97_unique.csv"), index=False)

# Full genome-wide rankings for all 7 targets
for target in TARGET_GENES:
    if target in correlations:
        corr = correlations[target]
        out = pd.DataFrame({
            'gene': corr.index,
            f'r_with_{target}': corr.values,
            'rank': range(1, len(corr) + 1),
        })
        out.to_csv(os.path.join(RESULTS_DIR, f"{target}_genome_wide_rankings.csv"),
                   index=False)

print(f"  Shared: {len(shared_st)} genes")
print(f"  SIGMAR1-unique: {len(sigmar1_unique)} genes")
print(f"  TMEM97-unique: {len(tmem97_unique)} genes")
print(f"  Genome-wide rankings: 7 files")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 16: FIGURE 1 — DIVERGENT NETWORKS                                 ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("GENERATING FIGURES")
print(f"{'='*70}\n")

fig = plt.figure(figsize=(18, 6))

# ── Panel A: Venn diagram ──
ax_a = fig.add_axes([0.02, 0.12, 0.28, 0.80])
ax_a.text(-0.05, 1.08, 'A', fontsize=20, fontweight='bold', va='top',
          transform=ax_a.transAxes)

n_s_only = len(sigmar1_unique)
n_t_only = len(tmem97_unique)
n_shared_venn = len(shared_st)

v = venn2(subsets=(n_s_only, n_t_only, n_shared_venn),
          set_labels=('SIGMAR1', 'TMEM97'), ax=ax_a)
v.get_patch_by_id('10').set_color('#1565C0')
v.get_patch_by_id('01').set_color('#E65100')
v.get_patch_by_id('11').set_color('#7B1FA2')
for patch in ['10', '01', '11']:
    v.get_patch_by_id(patch).set_alpha(0.7)
    v.get_patch_by_id(patch).set_edgecolor('#555')

j_val = n_shared_venn / (n_s_only + n_t_only + n_shared_venn)
ax_a.set_title(f'SIGMAR1–TMEM97 Network Overlap\n'
               f'Jaccard = {j_val:.3f}',
               fontsize=12, fontweight='bold')

# ── Panel B: Scatter plot of genome-wide rankings ──
ax_b = fig.add_axes([0.35, 0.12, 0.28, 0.80])
ax_b.text(-0.15, 1.08, 'B', fontsize=20, fontweight='bold', va='top',
          transform=ax_b.transAxes)

common_plot = sorted(set(correlations['SIGMAR1'].index) &
                     set(correlations['TMEM97'].index))
ranks_s = correlations['SIGMAR1'].reindex(common_plot).rank(ascending=False)
ranks_t = correlations['TMEM97'].reindex(common_plot).rank(ascending=False)

np.random.seed(42)
idx = np.random.choice(len(common_plot),
                       min(5000, len(common_plot)), replace=False)
ax_b.scatter(ranks_s.values[idx], ranks_t.values[idx],
            s=1, alpha=0.3, c='#555', rasterized=True)

rho_plot = st_row['spearman_rho'].values[0]
ax_b.set_xlabel('SIGMAR1 rank', fontsize=11)
ax_b.set_ylabel('TMEM97 rank', fontsize=11)
ax_b.set_title(f'Genome-wide rank correlation\nρ = {rho_plot:.3f}',
               fontsize=12, fontweight='bold')
ax_b.spines['top'].set_visible(False)
ax_b.spines['right'].set_visible(False)

# ── Panel C: Top 10 co-expression partners ──
ax_c = fig.add_axes([0.70, 0.12, 0.28, 0.80])
ax_c.text(-0.15, 1.08, 'C', fontsize=20, fontweight='bold', va='top',
          transform=ax_c.transAxes)

top10_s = correlations['SIGMAR1'].head(10)
top10_t = correlations['TMEM97'].head(10)

y_positions = np.arange(10)
bar_height = 0.35

bars1 = ax_c.barh(y_positions + bar_height/2, top10_s.values[::-1],
                   bar_height, color='#1565C0', alpha=0.85, label='SIGMAR1')
bars2 = ax_c.barh(y_positions - bar_height/2, top10_t.values[::-1],
                   bar_height, color='#E65100', alpha=0.85, label='TMEM97')

ax_c.set_yticks(y_positions)
s_labels = list(top10_s.index)[::-1]
t_labels = list(top10_t.index)[::-1]
combined_labels = [f'{s} / {t}' for s, t in zip(s_labels, t_labels)]
ax_c.set_yticklabels(combined_labels, fontsize=7)
ax_c.set_xlabel('Pearson r', fontsize=10)
ax_c.set_title('Top 10 Co-expression Partners', fontsize=12, fontweight='bold')
ax_c.legend(fontsize=9, loc='lower right')
ax_c.spines['top'].set_visible(False)
ax_c.spines['right'].set_visible(False)

fig.savefig(os.path.join(FIGURES_DIR, 'Figure1_Divergent_Networks.png'),
            dpi=300, bbox_inches='tight', facecolor='white')
fig.savefig(os.path.join(FIGURES_DIR, 'Figure1_Divergent_Networks.pdf'),
            dpi=300, bbox_inches='tight', facecolor='white')
plt.show()
print("✓ Figure 1 saved")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 17: FIGURE 2 — GO ENRICHMENT COMPARISON                           ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

def extract_top_terms(go_df, source='GO:BP', n=5):
    """Extract top n terms from gProfiler result DataFrame."""
    if go_df is None or len(go_df) == 0:
        return []
    subset = go_df[go_df['source'] == source].sort_values('p_value').head(n)
    terms = []
    for _, row in subset.iterrows():
        terms.append({
            'name': row['name'][:50],
            'p': row['p_value'],
            'nlp': -np.log10(max(row['p_value'], 1e-50)),
        })
    return terms

# Get terms
s_uniq_bp = extract_top_terms(go_sigmar1_unique, 'GO:BP', 5)
t_uniq_bp = extract_top_terms(go_tmem97_unique, 'GO:BP', 5)
shared_bp = extract_top_terms(go_shared, 'GO:BP', 5)

# Also Reactome for each
s_reac = extract_top_terms(go_sigmar1_full, 'REAC', 5)
t_reac = extract_top_terms(go_tmem97_full, 'REAC', 5)
s_kegg = extract_top_terms(go_sigmar1_full, 'KEGG', 5)
t_kegg = extract_top_terms(go_tmem97_full, 'KEGG', 5)

fig, axes = plt.subplots(2, 3, figsize=(22, 10))
plt.subplots_adjust(wspace=0.65, hspace=0.45)

panels = [
    (0, 0, f'A  SIGMAR1 unique ({len(sigmar1_unique)}) — GO:BP',
     s_uniq_bp, '#1565C0'),
    (0, 1, f'B  TMEM97 unique ({len(tmem97_unique)}) — GO:BP',
     t_uniq_bp, '#E65100'),
    (0, 2, f'C  Shared ({len(shared_st)}) — GO:BP',
     shared_bp, '#7B1FA2'),
    (1, 0, f'D  SIGMAR1 full — Reactome',
     s_reac, '#1565C0'),
    (1, 1, f'E  TMEM97 full — Reactome',
     t_reac, '#E65100'),
    (1, 2, f'F  SIGMAR1 full — KEGG',
     s_kegg, '#0D47A1'),
]

for row_idx, col_idx, title, terms, color in panels:
    ax = axes[row_idx, col_idx]
    if not terms:
        ax.text(0.5, 0.5, 'No significant terms',
                ha='center', va='center', fontsize=12,
                transform=ax.transAxes)
        ax.set_title(title, fontsize=10, fontweight='bold')
        continue

    names = [t['name'] for t in reversed(terms)]
    values = [t['nlp'] for t in reversed(terms)]
    ax.barh(range(len(names)), values, color=color, alpha=0.85, edgecolor='white')
    ax.set_yticks(range(len(names)))
    ax.set_yticklabels(names, fontsize=8)
    ax.set_xlabel('$-\\log_{10}(p)$', fontsize=10)
    ax.set_title(title, fontsize=10, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

fig.savefig(os.path.join(FIGURES_DIR, 'Figure2_GO_Enrichment.png'),
            dpi=300, bbox_inches='tight', facecolor='white')
fig.savefig(os.path.join(FIGURES_DIR, 'Figure2_GO_Enrichment.pdf'),
            dpi=300, bbox_inches='tight', facecolor='white')
plt.show()
print("✓ Figure 2 saved")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 18: FIGURE 3 — MULTI-REGION REPLICATION HEATMAP                   ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

fig, axes = plt.subplots(1, 2, figsize=(14, 5.5))
plt.subplots_adjust(wspace=0.35)

for idx, (target, title) in enumerate([('SIGMAR1', 'SIGMAR1'), ('TMEM97', 'TMEM97')]):
    ax = axes[idx]
    ax.text(-0.1, 1.08, chr(65 + idx), fontsize=20, fontweight='bold',
            va='top', transform=ax.transAxes)

    mat = cross_region_matrix[target]
    im = ax.imshow(mat, cmap='YlOrRd', vmin=0.75, vmax=1.0, aspect='equal')

    for i in range(5):
        for j in range(5):
            color = 'white' if mat[i, j] > 0.92 else 'black'
            ax.text(j, i, f'{mat[i, j]:.3f}', ha='center', va='center',
                   fontsize=9, fontweight='bold', color=color)

    ax.set_xticks(range(5))
    ax.set_yticks(range(5))
    ax.set_xticklabels(region_keys, fontsize=9, rotation=45, ha='right')
    ax.set_yticklabels(region_keys, fontsize=9)
    ax.set_title(f'{title} Cross-Region ρ', fontsize=12, fontweight='bold')
    cbar = plt.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    cbar.set_label('Spearman ρ', fontsize=9)

fig.savefig(os.path.join(FIGURES_DIR, 'Figure3_MultiRegion_Replication.png'),
            dpi=300, bbox_inches='tight', facecolor='white')
fig.savefig(os.path.join(FIGURES_DIR, 'Figure3_MultiRegion_Replication.pdf'),
            dpi=300, bbox_inches='tight', facecolor='white')
plt.show()
print("✓ Figure 3 saved")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 19: SUPPLEMENTARY TABLE S1 — GENOME-WIDE CORRELATION RANKINGS     ║
# ║                                                                         ║
# ║  16,224 genes × {Gene, SIGMAR1_r, SIGMAR1_Rank, TMEM97_r, TMEM97_Rank, ║
# ║  In_SIGMAR1_Top5%, In_TMEM97_Top5%, Classification}                     ║
# ║                                                                         ║
# ║  Saved as .xlsx (too large for Word table)                               ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("GENERATING SUPPLEMENTARY TABLE S1")
print("Genome-Wide Correlation Rankings")
print(f"{'='*70}\n")

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Build unified gene list
all_genes_union = sorted(
    set(correlations['SIGMAR1'].index) | set(correlations['TMEM97'].index)
)

s1_rows = []
for gene in all_genes_union:
    s1_r = correlations['SIGMAR1'].get(gene, np.nan)
    t_r = correlations['TMEM97'].get(gene, np.nan)

    # Ranks (1 = highest correlation)
    s1_rank = int((correlations['SIGMAR1'] > s1_r).sum() + 1) if not np.isnan(s1_r) else None
    t_rank = int((correlations['TMEM97'] > t_r).sum() + 1) if not np.isnan(t_r) else None

    in_s = gene in networks['SIGMAR1']
    in_t = gene in networks['TMEM97']

    if in_s and in_t:
        classification = 'Shared'
    elif in_s:
        classification = 'SIGMAR1-unique'
    elif in_t:
        classification = 'TMEM97-unique'
    else:
        classification = 'Neither'

    s1_rows.append({
        'Gene_Symbol': gene,
        'SIGMAR1_r': round(s1_r, 6) if not np.isnan(s1_r) else None,
        'SIGMAR1_Rank': s1_rank,
        'TMEM97_r': round(t_r, 6) if not np.isnan(t_r) else None,
        'TMEM97_Rank': t_rank,
        'In_SIGMAR1_Top5pct': 'Yes' if in_s else 'No',
        'In_TMEM97_Top5pct': 'Yes' if in_t else 'No',
        'Classification': classification,
    })

s1_df = pd.DataFrame(s1_rows)
# Sort by SIGMAR1 rank
s1_df = s1_df.sort_values('SIGMAR1_Rank', na_position='last').reset_index(drop=True)

# ── Write formatted .xlsx ──
wb = Workbook()
ws = wb.active
ws.title = "Table S1"

# Title rows
title_font = Font(name='Times New Roman', bold=True, size=12)
body_font = Font(name='Times New Roman', size=10)
header_font = Font(name='Times New Roman', bold=True, size=10, color='FFFFFF')
header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
shared_fill = PatternFill(start_color='E8D5F5', end_color='E8D5F5', fill_type='solid')
sigmar1_fill = PatternFill(start_color='D6EAF8', end_color='D6EAF8', fill_type='solid')
tmem97_fill = PatternFill(start_color='FDEBD0', end_color='FDEBD0', fill_type='solid')
thin_border = Border(
    left=Side(style='thin', color='CCCCCC'),
    right=Side(style='thin', color='CCCCCC'),
    top=Side(style='thin', color='CCCCCC'),
    bottom=Side(style='thin', color='CCCCCC'),
)

ws.merge_cells('A1:H1')
ws['A1'] = ("Table S1. Genome-wide Pearson correlation rankings for SIGMAR1 and "
            "TMEM97 in human frontal cortex (BA9).")
ws['A1'].font = title_font
ws['A1'].alignment = Alignment(wrap_text=True)
ws.row_dimensions[1].height = 35

ws.merge_cells('A2:H2')
ws['A2'] = (f"GTEx v8 RNA-seq, n = {n_samples_ba9} donors, "
            f"{len(s1_df)} protein-coding genes (median TPM ≥ 1.0). "
            f"Top 5% threshold: SIGMAR1 r ≥ {thresholds['SIGMAR1']:.3f}, "
            f"TMEM97 r ≥ {thresholds['TMEM97']:.3f}. "
            f"Classification: Shared = in both top 5% networks; "
            f"unique = in one network only.")
ws['A2'].font = Font(name='Times New Roman', size=9, italic=True)
ws['A2'].alignment = Alignment(wrap_text=True)
ws.row_dimensions[2].height = 45

# Column headers (row 3)
headers = ['Gene Symbol', 'SIGMAR1 r', 'SIGMAR1 Rank', 'TMEM97 r',
           'TMEM97 Rank', 'In SIGMAR1 Top 5%', 'In TMEM97 Top 5%',
           'Classification']
col_widths = [16, 13, 14, 13, 13, 18, 17, 18]

for col_idx, (h, w) in enumerate(zip(headers, col_widths), 1):
    cell = ws.cell(row=3, column=col_idx, value=h)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    cell.border = thin_border
    ws.column_dimensions[get_column_letter(col_idx)].width = w

ws.row_dimensions[3].height = 30

# Data rows (starting row 4)
for i, row_data in enumerate(s1_rows):
    row_num = i + 4
    vals = [
        row_data['Gene_Symbol'],
        row_data['SIGMAR1_r'],
        row_data['SIGMAR1_Rank'],
        row_data['TMEM97_r'],
        row_data['TMEM97_Rank'],
        row_data['In_SIGMAR1_Top5pct'],
        row_data['In_TMEM97_Top5pct'],
        row_data['Classification'],
    ]

    # Row fill based on classification
    if row_data['Classification'] == 'Shared':
        fill = shared_fill
    elif row_data['Classification'] == 'SIGMAR1-unique':
        fill = sigmar1_fill
    elif row_data['Classification'] == 'TMEM97-unique':
        fill = tmem97_fill
    else:
        fill = None

    for col_idx, val in enumerate(vals, 1):
        cell = ws.cell(row=row_num, column=col_idx, value=val)
        cell.font = body_font
        cell.border = thin_border
        if col_idx >= 2:
            cell.alignment = Alignment(horizontal='center')
        if fill:
            cell.fill = fill

    if row_num % 5000 == 0:
        print(f"  S1: wrote {row_num - 3:,} / {len(s1_rows):,} rows...")

# Freeze header row
ws.freeze_panes = 'A4'

# Auto-filter
ws.auto_filter.ref = f"A3:H{len(s1_rows) + 3}"

s1_path = os.path.join(SUPPL_DIR, "Table_S1_Genome_Wide_Rankings.xlsx")
wb.save(s1_path)
print(f"\n✓ Table S1 saved: {s1_path}")
print(f"  {len(s1_df)} genes, {len(s1_df[s1_df['Classification']=='Shared'])} shared, "
      f"{len(s1_df[s1_df['Classification']=='SIGMAR1-unique'])} SIGMAR1-unique, "
      f"{len(s1_df[s1_df['Classification']=='TMEM97-unique'])} TMEM97-unique")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 20: SUPPLEMENTARY TABLE S2 — gProfiler ENRICHMENT RESULTS         ║
# ║                                                                         ║
# ║  Multi-sheet .xlsx: one sheet per gene set query                         ║
# ║  (SIGMAR1 full, TMEM97 full, SIGMAR1-unique, TMEM97-unique, Shared)     ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("GENERATING SUPPLEMENTARY TABLE S2")
print("gProfiler Functional Enrichment Results")
print(f"{'='*70}\n")

wb2 = Workbook()
wb2.remove(wb2.active)  # Remove default sheet

gprofiler_sets = OrderedDict([
    ('SIGMAR1_full', go_sigmar1_full),
    ('TMEM97_full', go_tmem97_full),
    ('SIGMAR1_unique', go_sigmar1_unique),
    ('TMEM97_unique', go_tmem97_unique),
    ('Shared', go_shared),
])

# Columns to include from gProfiler output
gprofiler_cols = ['source', 'native', 'name', 'p_value', 'significant',
                  'term_size', 'query_size', 'intersection_size',
                  'effective_domain_size', 'precision', 'recall']

for sheet_name, df in gprofiler_sets.items():
    ws2 = wb2.create_sheet(title=sheet_name[:31])  # Excel 31-char limit

    # Title row
    ws2.merge_cells(f'A1:{get_column_letter(len(gprofiler_cols))}1')
    ws2['A1'] = (f"Table S2 — gProfiler enrichment: {sheet_name.replace('_', ' ')} "
                 f"(g:SCS threshold < 0.05, custom background = all expressed genes)")
    ws2['A1'].font = title_font
    ws2['A1'].alignment = Alignment(wrap_text=True)
    ws2.row_dimensions[1].height = 30

    if df is None or len(df) == 0:
        ws2['A2'] = "No significant terms identified."
        ws2['A2'].font = Font(name='Times New Roman', size=10, italic=True)
        print(f"  {sheet_name}: 0 terms (empty sheet)")
        continue

    # Filter and sort
    available_cols = [c for c in gprofiler_cols if c in df.columns]
    df_out = df[available_cols].sort_values('p_value').reset_index(drop=True)

    # Column headers (row 2)
    header_labels = {
        'source': 'Source', 'native': 'Term ID', 'name': 'Term Name',
        'p_value': 'Adj. p-value', 'significant': 'Significant',
        'term_size': 'Term Size', 'query_size': 'Query Size',
        'intersection_size': 'Overlap', 'effective_domain_size': 'Domain Size',
        'precision': 'Precision', 'recall': 'Recall',
    }

    for col_idx, col_name in enumerate(available_cols, 1):
        cell = ws2.cell(row=2, column=col_idx,
                        value=header_labels.get(col_name, col_name))
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', wrap_text=True)
        cell.border = thin_border

    # Data rows
    for row_idx, (_, row_data) in enumerate(df_out.iterrows(), 3):
        for col_idx, col_name in enumerate(available_cols, 1):
            val = row_data[col_name]
            # Format p-values
            if col_name == 'p_value' and isinstance(val, (float, np.floating)):
                cell = ws2.cell(row=row_idx, column=col_idx, value=float(val))
                cell.number_format = '0.00E+00'
            elif col_name in ('precision', 'recall') and isinstance(val, (float, np.floating)):
                cell = ws2.cell(row=row_idx, column=col_idx, value=float(val))
                cell.number_format = '0.0000'
            else:
                cell = ws2.cell(row=row_idx, column=col_idx, value=val)
            cell.font = body_font
            cell.border = thin_border
            if col_name != 'name':
                cell.alignment = Alignment(horizontal='center')

    # Column widths
    width_map = {'source': 10, 'native': 16, 'name': 45, 'p_value': 14,
                 'significant': 12, 'term_size': 12, 'query_size': 12,
                 'intersection_size': 10, 'effective_domain_size': 14,
                 'precision': 12, 'recall': 10}
    for col_idx, col_name in enumerate(available_cols, 1):
        ws2.column_dimensions[get_column_letter(col_idx)].width = width_map.get(col_name, 14)

    ws2.freeze_panes = 'A3'
    print(f"  {sheet_name}: {len(df_out)} terms")

s2_path = os.path.join(SUPPL_DIR, "Table_S2_gProfiler_Enrichment.xlsx")
wb2.save(s2_path)
print(f"\n✓ Table S2 saved: {s2_path}")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 21: SUPPLEMENTARY TABLE S3 — CUSTOM GENE SET ENRICHMENT (.docx)   ║
# ║                                                                         ║
# ║  Fisher's exact test results for 6 curated gene sets ×                  ║
# ║  SIGMAR1 and TMEM97 networks                                            ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("GENERATING SUPPLEMENTARY TABLE S3")
print("Custom Gene Set Enrichment (Word Document)")
print(f"{'='*70}\n")

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color_hex):
    """Apply shading to a Word table cell."""
    shading_elm = cell._element.get_or_add_tcPr().makeelement(
        qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): color_hex,
        }
    )
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, color='999999', size='4'):
    """Set thin borders on a table cell."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right'):
        element = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): size,
            qn('w:space'): '0',
            qn('w:color'): color,
        })
        borders.append(element)
    tcPr.append(borders)

doc_s3 = DocxDocument()

# Set default font
style = doc_s3.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# ── Title ──
title_para = doc_s3.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title_para.add_run(
    "Table S3. Custom gene set enrichment analysis for SIGMAR1 and TMEM97 "
    "top 5% co-expression networks in human frontal cortex (BA9)."
)
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

# ── Methods note ──
note_para = doc_s3.add_paragraph()
note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_note = note_para.add_run(
    "Fisher\u2019s exact test (one-sided, greater) against all expressed genes "
    f"(n = {gene_universe + 1}). Six curated gene sets tested against each "
    "receptor\u2019s top 5% co-expression network. Fold enrichment = observed / "
    "expected overlap. Genes listed are those present in both the custom set "
    "and the top 5% network."
)
run_note.font.size = Pt(9)
run_note.font.name = 'Times New Roman'
run_note.italic = True

doc_s3.add_paragraph()  # spacer

# ── Build table ──
# Columns: Gene Set | Set Size | Expressed | Target | In Network | Fold | p-value | Genes
col_headers = ['Gene Set', 'Set Size', 'Expressed',
               'Target', 'In Network', 'Fold Enrichment',
               'p-value', 'Genes in Network']
col_widths_inches = [1.6, 0.7, 0.8, 0.9, 0.8, 1.0, 1.0, 2.7]

# Collect all results into rows
s3_data_rows = []
for target in PRIMARY_TARGETS:
    for result in all_custom_results[target]:
        s3_data_rows.append({
            'gene_set': result['gene_set'],
            'set_size': len(custom_sets[result['gene_set']]),
            'expressed': result['expressed'],
            'target': target,
            'in_network': result['in_network'],
            'fold': result['fold_enrichment'],
            'p_value': result['p_value'],
            'genes': result['genes_found'],
        })

n_data_rows = len(s3_data_rows)
table_s3 = doc_s3.add_table(rows=1 + n_data_rows, cols=8)
table_s3.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for j, header_text in enumerate(col_headers):
    cell = table_s3.rows[0].cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header_text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    set_cell_shading(cell, '1F4E79')
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_borders(cell)

# Data rows
for i, row_data in enumerate(s3_data_rows):
    row_cells = table_s3.rows[i + 1].cells

    values = [
        row_data['gene_set'],
        str(row_data['set_size']),
        str(row_data['expressed']),
        row_data['target'],
        str(row_data['in_network']),
        f"{row_data['fold']:.1f}\u00d7",
        f"{row_data['p_value']:.2e}" if row_data['p_value'] < 0.001 else f"{row_data['p_value']:.3f}",
        row_data['genes'] if row_data['in_network'] > 0 else '\u2014',
    ]

    # Determine significance highlighting
    is_sig = row_data['p_value'] < 0.05

    for j, val in enumerate(values):
        cell = row_cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 7 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        if is_sig:
            run.bold = True
        set_cell_borders(cell)

        # Light green for significant rows
        if is_sig:
            set_cell_shading(cell, 'E8F5E9')

    if (i + 1) % 6 == 0 and i < n_data_rows - 1:
        # Visual separator between SIGMAR1 and TMEM97 blocks: slightly darker
        pass  # border handles it

# Set column widths
for j, w in enumerate(col_widths_inches):
    for row in table_s3.rows:
        row.cells[j].width = Inches(w)

# ── Footnote ──
doc_s3.add_paragraph()
fn_para = doc_s3.add_paragraph()
fn_note = fn_para.add_run(
    "Bold rows indicate statistically significant enrichment (p < 0.05). "
    "Gene set definitions: MAM-mitochondrial = mitochondria-associated membrane "
    "markers; Sigma receptor network = established sigma receptor interactors; "
    "ER stress/UPR = unfolded protein response pathway; Methylation pathway = "
    "one-carbon metabolism and SAM cycle; Vascular markers = endothelial-specific "
    "genes (negative control); Ribosome quality control = RQC pathway components "
    "including extraction machinery."
)
fn_note.font.size = Pt(8)
fn_note.font.name = 'Times New Roman'
fn_note.italic = True

s3_path = os.path.join(SUPPL_DIR, "Table_S3_Custom_Gene_Set_Enrichment.docx")
doc_s3.save(s3_path)
print(f"✓ Table S3 saved: {s3_path}")
print(f"  {n_data_rows} rows ({len(custom_sets)} gene sets × {len(PRIMARY_TARGETS)} targets)")

# Also save as xlsx for reviewer convenience
wb3 = Workbook()
ws3 = wb3.active
ws3.title = "Table S3"
ws3.merge_cells('A1:H1')
ws3['A1'] = ("Table S3. Custom gene set enrichment — Fisher's exact test, "
             "SIGMAR1 and TMEM97 top 5% networks.")
ws3['A1'].font = title_font

s3_headers = ['Gene Set', 'Set Size', 'Expressed in BA9', 'Target Network',
              'Genes in Network', 'Fold Enrichment', 'p-value', 'Genes Found']
for j, h in enumerate(s3_headers, 1):
    cell = ws3.cell(row=2, column=j, value=h)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal='center', wrap_text=True)
    cell.border = thin_border

for i, rd in enumerate(s3_data_rows, 3):
    ws3.cell(row=i, column=1, value=rd['gene_set']).font = body_font
    ws3.cell(row=i, column=2, value=rd['set_size']).font = body_font
    ws3.cell(row=i, column=3, value=rd['expressed']).font = body_font
    ws3.cell(row=i, column=4, value=rd['target']).font = body_font
    ws3.cell(row=i, column=5, value=rd['in_network']).font = body_font
    c_fold = ws3.cell(row=i, column=6, value=round(rd['fold'], 2))
    c_fold.font = body_font
    c_fold.number_format = '0.0'
    c_p = ws3.cell(row=i, column=7, value=float(rd['p_value']))
    c_p.font = body_font
    c_p.number_format = '0.00E+00'
    ws3.cell(row=i, column=8, value=rd['genes']).font = body_font
    for j in range(1, 9):
        ws3.cell(row=i, column=j).border = thin_border
        ws3.cell(row=i, column=j).alignment = Alignment(
            horizontal='center' if j != 8 else 'left')

for j, w in zip(range(1, 9), [22, 10, 14, 16, 14, 14, 14, 45]):
    ws3.column_dimensions[get_column_letter(j)].width = w

ws3.freeze_panes = 'A3'
s3_xlsx_path = os.path.join(SUPPL_DIR, "Table_S3_Custom_Gene_Set_Enrichment.xlsx")
wb3.save(s3_xlsx_path)
print(f"✓ Table S3 (xlsx) saved: {s3_xlsx_path}")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 22: SUPPLEMENTARY TABLE S4 — SHARED GENES (.docx + .xlsx)         ║
# ║                                                                         ║
# ║  All genes shared between SIGMAR1 and TMEM97 top 5% networks            ║
# ║  with correlation coefficients and ranks for both targets                ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print(f"\n{'='*70}")
print("GENERATING SUPPLEMENTARY TABLE S4")
print("Shared Gene Characterization")
print(f"{'='*70}\n")

# ── Build S4 data ──
s4_rows = []
for gene in sorted(shared_st):
    s_r = correlations['SIGMAR1'].get(gene, np.nan)
    t_r = correlations['TMEM97'].get(gene, np.nan)
    s_rank = int((correlations['SIGMAR1'] > s_r).sum() + 1) if not np.isnan(s_r) else None
    t_rank = int((correlations['TMEM97'] > t_r).sum() + 1) if not np.isnan(t_r) else None
    s4_rows.append({
        'gene': gene,
        'sigmar1_r': round(float(s_r), 4) if not np.isnan(s_r) else None,
        'sigmar1_rank': s_rank,
        'tmem97_r': round(float(t_r), 4) if not np.isnan(t_r) else None,
        'tmem97_rank': t_rank,
        'rank_diff': abs(s_rank - t_rank) if s_rank and t_rank else None,
    })

# Sort by mean rank (most correlated with both)
for row in s4_rows:
    if row['sigmar1_rank'] and row['tmem97_rank']:
        row['mean_rank'] = (row['sigmar1_rank'] + row['tmem97_rank']) / 2
    else:
        row['mean_rank'] = 99999
s4_rows.sort(key=lambda x: x['mean_rank'])

# ── Word document ──
doc_s4 = DocxDocument()
style = doc_s4.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)

# Title
tp = doc_s4.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = tp.add_run(
    f"Table S4. Genes shared between SIGMAR1 and TMEM97 top 5% co-expression "
    f"networks (n = {len(shared_st)}) in human frontal cortex (BA9)."
)
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

# Methods note
np2 = doc_s4.add_paragraph()
rn = np2.add_run(
    f"Pearson correlation coefficients computed from GTEx v8 RNA-seq "
    f"(n = {n_samples_ba9} donors, log\u2082(TPM+1) transformed). "
    f"Rank = position among {gene_universe} non-target genes. "
    f"Sorted by mean rank across both targets."
)
rn.font.size = Pt(9)
rn.font.name = 'Times New Roman'
rn.italic = True

doc_s4.add_paragraph()

# Table
s4_headers = ['Gene Symbol', 'SIGMAR1 r', 'SIGMAR1 Rank',
              'TMEM97 r', 'TMEM97 Rank', 'Rank Difference']
s4_widths = [1.5, 1.2, 1.2, 1.2, 1.2, 1.3]

table_s4 = doc_s4.add_table(rows=1 + len(s4_rows), cols=6)
table_s4.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for j, h in enumerate(s4_headers):
    cell = table_s4.rows[0].cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '1F4E79')
    set_cell_borders(cell)

# Data
for i, rd in enumerate(s4_rows):
    cells = table_s4.rows[i + 1].cells
    vals = [
        rd['gene'],
        f"{rd['sigmar1_r']:.4f}" if rd['sigmar1_r'] is not None else '\u2014',
        str(rd['sigmar1_rank']) if rd['sigmar1_rank'] is not None else '\u2014',
        f"{rd['tmem97_r']:.4f}" if rd['tmem97_r'] is not None else '\u2014',
        str(rd['tmem97_rank']) if rd['tmem97_rank'] is not None else '\u2014',
        str(rd['rank_diff']) if rd['rank_diff'] is not None else '\u2014',
    ]
    for j, val in enumerate(vals):
        cell = cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        set_cell_borders(cell)

        # Alternating row shading
        if i % 2 == 1:
            set_cell_shading(cell, 'F2F2F2')

for j, w in enumerate(s4_widths):
    for row in table_s4.rows:
        row.cells[j].width = Inches(w)

# Footnote
doc_s4.add_paragraph()
fn = doc_s4.add_paragraph()
fn_run = fn.add_run(
    f"Rank difference = |SIGMAR1 rank \u2212 TMEM97 rank|. "
    f"Low rank difference indicates genes that occupy similar positions in both "
    f"co-expression hierarchies. "
    f"SIGMAR1 top 5% threshold: r \u2265 {thresholds['SIGMAR1']:.3f}; "
    f"TMEM97 top 5% threshold: r \u2265 {thresholds['TMEM97']:.3f}."
)
fn_run.font.size = Pt(8)
fn_run.font.name = 'Times New Roman'
fn_run.italic = True

s4_docx_path = os.path.join(SUPPL_DIR, "Table_S4_Shared_Genes.docx")
doc_s4.save(s4_docx_path)
print(f"✓ Table S4 (docx) saved: {s4_docx_path}")
print(f"  {len(s4_rows)} shared genes")

# ── Also save as xlsx ──
wb4 = Workbook()
ws4 = wb4.active
ws4.title = "Table S4"

ws4.merge_cells('A1:F1')
ws4['A1'] = (f"Table S4. Shared genes between SIGMAR1 and TMEM97 top 5% "
             f"co-expression networks (n = {len(shared_st)}).")
ws4['A1'].font = title_font
ws4['A1'].alignment = Alignment(wrap_text=True)

for j, h in enumerate(s4_headers, 1):
    cell = ws4.cell(row=2, column=j, value=h)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal='center', wrap_text=True)
    cell.border = thin_border

for i, rd in enumerate(s4_rows, 3):
    ws4.cell(row=i, column=1, value=rd['gene']).font = body_font
    c_sr = ws4.cell(row=i, column=2, value=rd['sigmar1_r'])
    c_sr.font = body_font
    c_sr.number_format = '0.0000'
    ws4.cell(row=i, column=3, value=rd['sigmar1_rank']).font = body_font
    c_tr = ws4.cell(row=i, column=4, value=rd['tmem97_r'])
    c_tr.font = body_font
    c_tr.number_format = '0.0000'
    ws4.cell(row=i, column=5, value=rd['tmem97_rank']).font = body_font
    ws4.cell(row=i, column=6, value=rd['rank_diff']).font = body_font
    for j in range(1, 7):
        ws4.cell(row=i, column=j).border = thin_border
        ws4.cell(row=i, column=j).alignment = Alignment(horizontal='center')

for j, w in zip(range(1, 7), [16, 13, 14, 13, 14, 16]):
    ws4.column_dimensions[get_column_letter(j)].width = w

ws4.freeze_panes = 'A3'
s4_xlsx_path = os.path.join(SUPPL_DIR, "Table_S4_Shared_Genes.xlsx")
wb4.save(s4_xlsx_path)
print(f"✓ Table S4 (xlsx) saved: {s4_xlsx_path}")

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  CELL 23: FINAL SUMMARY — ALL MANUSCRIPT NUMBERS                        ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

print("\n" + "=" * 70)
print("DEFINITIVE MANUSCRIPT NUMBERS — MS3 SIGMA DIVERGENCE")
print("Target: Journal of Neurochemistry (JNC)")
print("=" * 70)

print(f"\n─── Basic Parameters ───")
print(f"  Gene universe (with target): {gene_universe + 1}")
print(f"  Non-target genes: {gene_universe}")
print(f"  Samples (BA9): {n_samples_ba9}")

print(f"\n─── SIGMAR1 Network ───")
print(f"  Top 5% count: {len(networks['SIGMAR1'])}")
print(f"  Threshold: r ≥ {thresholds['SIGMAR1']:.3f}")
print(f"  Top partner: {correlations['SIGMAR1'].idxmax()} "
      f"(r = {correlations['SIGMAR1'].max():.3f})")

print(f"\n─── TMEM97 Network ───")
print(f"  Top 5% count: {len(networks['TMEM97'])}")
print(f"  Threshold: r ≥ {thresholds['TMEM97']:.3f}")
print(f"  Top partner: {correlations['TMEM97'].idxmax()} "
      f"(r = {correlations['TMEM97'].max():.3f})")

print(f"\n─── SIGMAR1–TMEM97 Divergence ───")
j_val = st_row['jaccard'].values[0]
shared_val = int(st_row['shared'].values[0])
or_val = st_row['fisher_or'].values[0]
p_val = st_row['fisher_p'].values[0]
rho_val = st_row['spearman_rho'].values[0]
print(f"  Shared genes: {shared_val}")
print(f"  Jaccard: {j_val:.3f}")
print(f"  Fisher OR: {or_val:.2f}")
print(f"  Fisher p: {p_val:.2e}")
print(f"  Spearman ρ: {rho_val:.3f}")
print(f"  SIGMAR1 unique: {len(sigmar1_unique)}")
print(f"  TMEM97 unique: {len(tmem97_unique)}")

print(f"\n─── SIGMAR1–LTN1 Overlap ───")
if len(sl_row) > 0:
    print(f"  Shared: {int(sl_row['shared'].values[0])}")
    print(f"  Jaccard: {sl_row['jaccard'].values[0]:.3f}")
    print(f"  Spearman ρ: {sl_row['spearman_rho'].values[0]:.3f}")

print(f"\n─── 7-Gene Panel Context ───")
for _, row in comp_df.iterrows():
    g1, g2 = row['gene1'], row['gene2']
    print(f"  {g1:8s}–{g2:8s}: J={row['jaccard']:.3f}, "
          f"shared={int(row['shared']):4d}, ρ={row['spearman_rho']:.3f}")

print(f"\n─── Cross-Region Replication ───")
for target in PRIMARY_TARGETS:
    vals = cross_region_matrix[target][np.triu_indices(5, k=1)]
    print(f"  {target}: ρ = {vals.min():.3f} – {vals.max():.3f}")

print(f"\n─── Sensitivity Analyses ───")
for target in PRIMARY_TARGETS:
    if target in ct_results:
        common = sorted(set(ct_results[target]['corr'].index) &
                        set(correlations[target].index))
        rho_ct, _ = spearmanr(
            ct_results[target]['corr'].reindex(common).values,
            correlations[target].reindex(common).values)
        print(f"  {target} cell-type adjusted rank ρ: {rho_ct:.4f}")
    if target in agesex_results:
        common = sorted(set(agesex_results[target]['corr'].index) &
                        set(correlations[target].index))
        rho_as, _ = spearmanr(
            agesex_results[target]['corr'].reindex(common).values,
            correlations[target].reindex(common).values)
        print(f"  {target} covariate-adjusted rank ρ: {rho_as:.4f}")

print(f"\n─── Output Files ───")
print(f"  Results:  {RESULTS_DIR}")
print(f"  Figures:  {FIGURES_DIR}")
print(f"  Table S1: {s1_path}")
print(f"  Table S2: {s2_path}")
print(f"  Table S3: {s3_path}")
print(f"  Table S3: {s3_xlsx_path}")
print(f"  Table S4: {s4_docx_path}")
print(f"  Table S4: {s4_xlsx_path}")

print("\n" + "=" * 70)
print("✓ PIPELINE COMPLETE — All results saved to Google Drive")
print(f"  {DRIVE_BASE}")
print("=" * 70)